"""试卷 Markdown、DOCX 和 PDF 导出 worker。

Java 创建导出任务并管理文件存储；本模块只负责把 Java 传入的试卷快照渲染为具体文件。
"""

import copy
import tempfile

from app.worker_base import *
from app.question_markdown import *
from app.import_services import *
from app.question_boundary import strip_sub_label

SUPERSCRIPT_DIGITS = str.maketrans("0123456789+-=()", "⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻⁼⁽⁾")
SUBSCRIPT_DIGITS = str.maketrans("0123456789+-=()", "₀₁₂₃₄₅₆₇₈₉₊₋₌₍₎")
LATEX_SYMBOLS = {
    r"\pi": "π",
    r"\infty": "∞",
    r"\angle": "∠",
    r"\circ": "°",
    r"\times": "×",
    r"\div": "÷",
    r"\cdot": "·",
    r"\leq": "≤",
    r"\le": "≤",
    r"\geq": "≥",
    r"\ge": "≥",
    r"\neq": "≠",
    r"\ne": "≠",
    r"\approx": "≈",
    r"\pm": "±",
    r"\alpha": "α",
    r"\beta": "β",
    r"\gamma": "γ",
    r"\theta": "θ",
    r"\Delta": "Δ",
}

def strip_markdown_image_lines(markdown: str) -> str:
    """移除 Markdown 中的图片行。"""
    return "\n".join(line for line in str(markdown or "").splitlines() if not re.search(r"!\[[^\]]*\]\([^)]+\)", line))


def replace_latex_groups(text: str) -> str:
    """替换 LaTeX 分组语法以适配导出。"""
    current = text
    fraction_re = re.compile(r"\\frac\s*\{([^{}]+)\}\s*\{([^{}]+)\}")
    sqrt_re = re.compile(r"\\sqrt\s*\{([^{}]+)\}")
    nth_root_re = re.compile(r"\\sqrt\s*\[([^][]+)\]\s*\{([^{}]+)\}")
    power_group_re = re.compile(r"\^\{([^{}]+)\}")
    sub_group_re = re.compile(r"_\{([^{}]+)\}")

    def render_fraction(match: re.Match[str]) -> str:
        """执行 render fraction 逻辑。"""
        numerator = match.group(1).strip()
        denominator = match.group(2).strip()
        if re.fullmatch(r"[\w+\-π∞]+", numerator) and re.fullmatch(r"[\w+\-π∞]+", denominator):
            return f"{numerator}/{denominator}"
        return f"({numerator})/({denominator})"

    def render_nth_root(match: re.Match[str]) -> str:
        """执行 render nth root 逻辑。"""
        degree = match.group(1).strip()
        radicand = match.group(2).strip()
        return f"{degree}√{radicand}"

    for _ in range(12):
        next_text = fraction_re.sub(render_fraction, current)
        next_text = nth_root_re.sub(render_nth_root, next_text)
        next_text = sqrt_re.sub(lambda match: f"√{match.group(1)}", next_text)
        next_text = power_group_re.sub(lambda match: f"^{match.group(1).translate(SUPERSCRIPT_DIGITS)}", next_text)
        next_text = sub_group_re.sub(lambda match: f"_{match.group(1).translate(SUBSCRIPT_DIGITS)}", next_text)
        if next_text == current:
            return next_text
        current = next_text
    return current


def render_latex_inline(content: str) -> str:
    """渲染内联 LaTeX 内容为导出友好文本。"""
    text = replace_latex_groups(content)
    text = re.sub(r"\\(?:left|right|displaystyle|textstyle|big|Big|bigg|Bigg)\s*", "", text)
    text = re.sub(r"\\begin\{array\}\s*\{[^{}]*\}", "", text)
    text = re.sub(r"\\end\{array\}", "", text)
    text = text.replace(r"\\", "；")
    text = text.replace(r"\{", "{").replace(r"\}", "}")
    for source, target in LATEX_SYMBOLS.items():
        text = re.sub(re.escape(source) + r"(?![a-zA-Z])", target, text)
    text = re.sub(r"\\text\s*\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\mathrm\s*\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\mathbf\s*\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\[a-zA-Z]+", "", text)
    text = text.replace("{", "").replace("}", "")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def render_markdown_for_export(markdown: str) -> str:
    """将 Markdown 转换为导出友好的纯文本。"""
    text = strip_markdown_image_lines(markdown)
    text = re.sub(r"\\begin\{tasks\}(?:\([^)]+\))?", "", text)
    text = re.sub(r"\\end\{tasks\}", "", text)
    text = re.sub(r"\\task\s*", "• ", text)
    text = re.sub(r"\$\$(.*?)\$\$", lambda match: render_latex_inline(match.group(1)), text, flags=re.S)
    text = re.sub(r"\$(.*?)\$", lambda match: render_latex_inline(match.group(1)), text, flags=re.S)
    text = re.sub(r"`([^`]*)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"__([^_]+)__", r"\1", text)
    text = re.sub(r"^\s{0,3}#{1,6}\s*", "", text, flags=re.M)
    text = re.sub(r"^\s{0,3}>\s?", "", text, flags=re.M)
    text = re.sub(r"^\s*[-*+]\s+", "", text, flags=re.M)
    text = re.sub(r"(?<!\*)\*(?!\*)([^*\n]+)(?<!\*)\*(?!\*)", r"\1", text)
    text = re.sub(r"(?<!_)_(?!_)([^_\n]+)(?<!_)_(?!_)", r"\1", text)
    return "\n".join(line.strip() for line in text.splitlines() if line.strip())


def paper_header_lines(paper: dict[str, Any], questions: list[dict[str, Any]]) -> list[str]:
    """生成试卷导出页眉信息。"""
    header = paper.get("header") or {}
    total_score = sum(float(question.get("score", 0) or 0) for question in questions)
    lines = [str(paper.get("title") or "组卷")]
    if header.get("subtitle"):
        lines.append(str(header["subtitle"]))
    meta_parts = []
    if header.get("school"):
        meta_parts.append(f"学校：{header['school']}")
    if header.get("duration"):
        meta_parts.append(f"考试时长：{header['duration']}")
    meta_parts.append(f"满分：{total_score:g} 分")
    meta_parts.append(f"题量：{len(questions)} 题")
    if meta_parts:
        lines.append("　".join(meta_parts))
    lines.append("姓名：__________　班级：__________　考号：__________")
    if header.get("instructions"):
        lines.append(f"考生须知：{render_markdown_for_export(str(header['instructions']))}")
    return lines


def question_sub_questions(question: dict[str, Any]) -> list[dict[str, Any]]:
    """读取题目小问，兼容 subQuestions 和 children。"""
    for key in ("subQuestions", "children"):
        value = question.get(key)
        if isinstance(value, list):
            return [item for item in value if isinstance(item, dict)]
    return []


def selected_sub_questions(question: dict[str, Any], sub_selections: dict[str, Any]) -> list[dict[str, Any]]:
    """按试卷 subSelections 过滤小问；缺失、为空或无交集时视为全选。"""
    subs = question_sub_questions(question)
    if not subs:
        return []
    selected_ids = sub_selections.get(str(question.get("id") or ""))
    if not isinstance(selected_ids, list) or not selected_ids:
        return subs
    selected_id_set = {str(item) for item in selected_ids}
    selected = [sub for sub in subs if str(sub.get("id") or "") in selected_id_set]
    return selected if selected else subs


def questions_for_export(paper: dict[str, Any], questions: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """复制导出题目并写入按试卷选择过滤后的小问。"""
    raw_sub_selections = paper.get("subSelections")
    sub_selections = raw_sub_selections if isinstance(raw_sub_selections, dict) else {}
    result: list[dict[str, Any]] = []
    for question in questions:
        if not isinstance(question, dict):
            continue
        item = copy.deepcopy(question)
        total_sub_count = len(question_sub_questions(item))
        subs = selected_sub_questions(item, sub_selections)
        if subs:
            item["_subQuestionTotal"] = total_sub_count
            item["subQuestions"] = subs
            item["children"] = subs
        result.append(item)
    return result


def effective_question_type(question: dict[str, Any], fallback_type: Any = None) -> str:
    """读取题型；小问题型缺失或 unknown 时继承父题题型。"""
    value = str(question.get("type") or "").strip()
    if value and value != "unknown":
        return value
    return str(fallback_type or value or "unknown").strip() or "unknown"


def numeric_score(value: Any) -> float:
    """读取分值，非法值按 0 处理。"""
    try:
        return float(value or 0)
    except (TypeError, ValueError):
        return 0.0


def answer_space_height(question: dict[str, Any], *, fallback_type: Any = None, is_sub_question: bool = False) -> float:
    """计算 PDF 中解答题自动预留的作答空间高度。"""
    if effective_question_type(question, fallback_type) != "solution":
        return 0
    score = numeric_score(question.get("score"))
    if is_sub_question:
        base_height = 72
        return max(base_height, min(150, base_height + score * 7))
    base_height = 118
    return max(base_height, min(260, base_height + score * 10))


def sub_question_markdown(sub_question: dict[str, Any], fallback_index: int) -> tuple[str, str, list[dict[str, Any]]]:
    """返回小问标签、题干 Markdown 和选项。"""
    label = str(sub_question.get("label") or f"({fallback_index})").strip()
    raw_markdown = str(sub_question.get("manualMarkdown") or sub_question.get("stemMarkdown") or sub_question.get("stem") or "")
    stem_markdown, task_options = split_tasks_options(raw_markdown)
    stem_markdown = strip_sub_label(stem_markdown, label)
    options = sub_question.get("options") or task_options
    return label, stem_markdown, options


def render_question_lines(question: dict[str, Any], number: int, include_answer: bool) -> list[str]:
    """生成单题导出的多行文本。"""
    raw_markdown = question.get("manualMarkdown") or question.get("stemMarkdown") or ""
    stem_markdown, task_options = split_tasks_options(str(raw_markdown))
    lines = [f"{number}. {render_markdown_for_export(stem_markdown)}"]
    options = question.get("options") or task_options
    for option in options:
        lines.append(f"{option.get('label')}. {render_markdown_for_export(str(option.get('content') or ''))}")
    score = question.get("score")
    if score is not None:
        lines.append(f"（本题 {float(score or 0):g} 分）")
    sub_questions = question_sub_questions(question)
    for sub_index, sub_question in enumerate(sub_questions, start=1):
        label, sub_stem_markdown, sub_options = sub_question_markdown(sub_question, sub_index)
        rendered_stem = render_markdown_for_export(sub_stem_markdown)
        lines.append(f"{label} {rendered_stem}".rstrip())
        for option in sub_options:
            lines.append(f"  {option.get('label')}. {render_markdown_for_export(str(option.get('content') or ''))}")
        if include_answer:
            if sub_question.get("answer"):
                lines.append(f"{label} 答案：{render_markdown_for_export(str(sub_question.get('answer')))}")
            if sub_question.get("analysis"):
                lines.append(f"{label} 解析：{render_markdown_for_export(str(sub_question.get('analysis')))}")
    if include_answer:
        if not sub_questions and question.get("answer"):
            lines.append(f"答案：{render_markdown_for_export(str(question.get('answer')))}")
        if not sub_questions and question.get("analysis"):
            lines.append(f"解析：{render_markdown_for_export(str(question.get('analysis')))}")
    return lines


def render_question_plain(question: dict[str, Any], number: int, include_answer: bool) -> str:
    """生成单题导出的纯文本。"""
    return "\n".join(render_question_lines(question, number, include_answer))


def safe_resolved_child(root: Path, *parts: str) -> Path | None:
    """在根目录内安全解析子路径。"""
    try:
        target = root.joinpath(*parts).resolve()
        resolved_root = root.resolve()
        if str(target).startswith(str(resolved_root)) and target.exists() and target.is_file():
            return target
    except Exception:
        return None
    return None


def resolve_export_image_path(image: dict[str, Any]) -> Path | None:
    """解析导出时可读取的题图路径。"""
    if not isinstance(image, dict):
        return None
    raw_values = [str(image.get(key) or "").strip() for key in ("url", "path", "name")]
    url = raw_values[0]
    path = raw_values[1]

    ocr_match = re.search(r"/api/ocr/jobs/([^/]+)/files/(.+)$", url)
    if ocr_match:
        target = safe_resolved_child(OUTPUT_ROOT / ocr_match.group(1), ocr_match.group(2))
        if target:
            return target

    upload_match = re.search(r"/api/import-tasks/([^/]+)/questions/([^/]+)/images/([^/?#]+)", url)
    if upload_match:
        target = import_question_image_file(upload_match.group(1), upload_match.group(2), safe_filename(upload_match.group(3))).resolve()
        upload_root = IMPORT_UPLOAD_ROOT.resolve()
        if str(target).startswith(str(upload_root)) and target.exists() and target.is_file():
            return target

    upload_path_match = re.search(r"question_uploads/([^/]+)/([^/]+)/([^/?#]+)", path)
    if upload_path_match:
        target = import_question_image_file(upload_path_match.group(1), upload_path_match.group(2), safe_filename(upload_path_match.group(3))).resolve()
        upload_root = IMPORT_UPLOAD_ROOT.resolve()
        if str(target).startswith(str(upload_root)) and target.exists() and target.is_file():
            return target

    normalized_path = normalize_asset_path(path)
    if normalized_path:
        direct_candidates = [
            safe_resolved_child(OUTPUT_ROOT, normalized_path),
            safe_resolved_child(IMPORT_UPLOAD_ROOT, normalized_path),
            safe_resolved_child(STORAGE_ROOT, normalized_path),
        ]
        for candidate in direct_candidates:
            if candidate:
                return candidate
        filename = Path(normalized_path).name
        if filename:
            for root in (OUTPUT_ROOT, IMPORT_UPLOAD_ROOT):
                for candidate in root.rglob(filename):
                    if candidate.is_file() and normalize_asset_path(candidate.as_posix()).endswith(normalized_path):
                        return candidate
            for root in (OUTPUT_ROOT, IMPORT_UPLOAD_ROOT):
                found = next((candidate for candidate in root.rglob(filename) if candidate.is_file()), None)
                if found:
                    return found
    return None


def export_question_image_paths(question: dict[str, Any], include_sub_questions: bool = True) -> list[Path]:
    """收集题目导出需要插入的图片路径。"""
    paths: list[Path] = []
    seen: set[str] = set()
    images = [*normalize_question_images(question.get("images", []))]
    if include_sub_questions:
        for sub_question in question_sub_questions(question):
            images.extend(normalize_question_images(sub_question.get("images", [])))
    for image in images:
        path = resolve_export_image_path(image)
        if not path:
            continue
        key = path.resolve().as_posix()
        if key in seen:
            continue
        seen.add(key)
        paths.append(path)
    return paths


def export_paper_docx_legacy(paper: dict[str, Any], questions: list[dict[str, Any]], include_answer: bool) -> Path:
    """使用 python-docx 旧路径导出 DOCX。"""
    questions = questions_for_export(paper, questions)
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt

    export_path = EXPORT_ROOT / f"{paper['id']}.docx"
    document = Document()

    normal_style = document.styles["Normal"]
    normal_style.font.name = "宋体"
    normal_style.font.size = Pt(11)
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    def add_text_paragraph(text: str, *, alignment: Any = WD_ALIGN_PARAGRAPH.LEFT, size: int = 11, space_before: int = 0, space_after: int = 4) -> None:
        """执行 add text paragraph 逻辑。"""
        paragraph = document.add_paragraph()
        paragraph.alignment = alignment
        paragraph.paragraph_format.space_before = Pt(space_before)
        paragraph.paragraph_format.space_after = Pt(space_after)
        paragraph.paragraph_format.line_spacing = 1.35
        run = paragraph.add_run(text)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(size)

    def add_image_paragraph(image_path: Path) -> None:
        """执行 add image paragraph 逻辑。"""
        try:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.add_run().add_picture(str(image_path), width=Inches(4.8))
        except Exception:
            add_text_paragraph(f"[题图无法导出：{image_path.name}]", size=10)

    header_lines = paper_header_lines(paper, questions)
    add_text_paragraph(header_lines[0], alignment=WD_ALIGN_PARAGRAPH.CENTER, size=14, space_after=6)
    for line in header_lines[1:]:
        alignment = WD_ALIGN_PARAGRAPH.CENTER if "：" not in line or line.startswith(("学校：", "姓名：")) else WD_ALIGN_PARAGRAPH.LEFT
        add_text_paragraph(line, alignment=alignment)
    add_text_paragraph("", space_after=2)
    for index, question in enumerate(questions, start=1):
        image_paths = export_question_image_paths(question)
        for line_index, line in enumerate(render_question_lines(question, index, include_answer)):
            add_text_paragraph(line, space_before=6 if line_index == 0 else 0)
            if line_index == 0:
                for image_path in image_paths:
                    add_image_paragraph(image_path)
        add_text_paragraph("", space_after=2)
    document.save(export_path)
    return export_path


def wrap_pdf_line(text: str, font_name: str, font_size: int, max_width: float) -> list[str]:
    """按 PDF 可用宽度拆分文本行。"""
    from reportlab.pdfbase.pdfmetrics import stringWidth

    lines: list[str] = []
    current = ""
    for char in text:
        next_text = current + char
        if current and stringWidth(next_text, font_name, font_size) > max_width:
            lines.append(current)
            current = char
        else:
            current = next_text
    if current:
        lines.append(current)
    return lines or [""]


def export_paper_pdf_legacy(paper: dict[str, Any], questions: list[dict[str, Any]], include_answer: bool) -> Path:
    """使用 reportlab 按前端预览样式导出 PDF。"""
    questions = questions_for_export(paper, questions)
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.utils import ImageReader
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.pdfgen import canvas
    from PIL import Image

    font_name = "Helvetica"
    regular_font_candidates = [
        Path("/System/Library/Fonts/Supplemental/Arial Unicode.ttf"),
        Path("/System/Library/Fonts/Hiragino Sans GB.ttc"),
        Path("/System/Library/Fonts/STHeiti Medium.ttc"),
        Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"),
        Path("/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc"),
    ]
    for font_path in regular_font_candidates:
        if not font_path.exists():
            continue
        try:
            pdfmetrics.registerFont(TTFont("ExamRegular", str(font_path), subfontIndex=0))
            font_name = "ExamRegular"
            break
        except Exception:
            continue
    if font_name == "Helvetica":
        try:
            pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
            font_name = "STSong-Light"
        except Exception:
            font_name = "Helvetica"
    export_path = EXPORT_ROOT / f"{paper['id']}.pdf"
    pdf = canvas.Canvas(str(export_path), pagesize=A4)
    width, height = A4
    margin_x = 64
    y = height - 48
    max_width = width - margin_x * 2
    title_size = 18
    body_size = 10.5
    small_size = 8.5
    line_height = 16
    page_bottom = 48

    ink = colors.HexColor("#0f172a")
    muted = colors.HexColor("#64748b")
    border = colors.HexColor("#dbe3ef")
    card_fill = colors.HexColor("#f8fafc")
    badge_fill = colors.HexColor("#eff6ff")
    badge_text = colors.HexColor("#0369a1")

    def type_label(value: Any) -> str:
        mapping = {
            "choice": "选择题",
            "fill_blank": "填空题",
            "solution": "解答题",
            "unknown": "未知",
        }
        return mapping.get(str(value or "unknown"), str(value or "未知"))

    def ensure_space(required: float = 16) -> None:
        """执行 ensure space 逻辑。"""
        nonlocal y
        if y < page_bottom + required:
            pdf.showPage()
            y = height - 48
            pdf.setFillColor(ink)
            pdf.setFont(font_name, body_size)

    def draw_text_lines(lines: list[str], x: float, available_width: float, *, font_size: float = body_size, color: Any = ink, leading: float | None = None) -> None:
        """绘制自动换行文本。"""
        nonlocal y
        leading = leading or font_size + 5
        pdf.setFillColor(color)
        pdf.setFont(font_name, font_size)
        for line in lines:
            for wrapped in wrap_pdf_line(str(line or ""), font_name, int(font_size), available_width):
                ensure_space(leading + 2)
                pdf.drawString(x, y, wrapped)
                y -= leading

    def markdown_lines(markdown: Any) -> list[str]:
        text = render_markdown_for_export(str(markdown or ""))
        return [line for line in text.splitlines() if line.strip()]

    def draw_badge(text: str, x: float, baseline_y: float) -> float:
        """绘制预览风格标签并返回下一个 x。"""
        from reportlab.pdfbase.pdfmetrics import stringWidth

        label = str(text or "").strip()
        if not label:
            return x
        padding_x = 7
        badge_h = 15
        badge_w = stringWidth(label, font_name, small_size) + padding_x * 2
        pdf.setStrokeColor(colors.HexColor("#bfdbfe"))
        pdf.setFillColor(badge_fill)
        pdf.roundRect(x, baseline_y - 3, badge_w, badge_h, 7, stroke=1, fill=1)
        pdf.setFillColor(badge_text)
        pdf.setFont(font_name, small_size)
        pdf.drawString(x + padding_x, baseline_y + 1, label)
        return x + badge_w + 6

    def option_lines(options: list[dict[str, Any]]) -> list[tuple[str, str]]:
        result: list[tuple[str, str]] = []
        for option in options:
            if not isinstance(option, dict):
                continue
            label = str(option.get("label") or "").strip()
            content = render_markdown_for_export(str(option.get("content") or option.get("contentMarkdown") or ""))
            if label or content:
                result.append((label, content))
        return result

    def draw_options(options: list[dict[str, Any]]) -> None:
        """按预览样式两列绘制选择题选项。"""
        nonlocal y
        items = option_lines(options)
        if not items:
            return
        gap = 8
        col_w = (max_width - gap) / 2
        box_h = 26
        for row_start in range(0, len(items), 2):
            row = items[row_start : row_start + 2]
            ensure_space(box_h + 8)
            row_y = y - box_h
            for col, (label, content) in enumerate(row):
                x = margin_x + col * (col_w + gap)
                pdf.setFillColor(colors.white)
                pdf.setStrokeColor(border)
                pdf.roundRect(x, row_y, col_w, box_h, 8, stroke=1, fill=1)
                pdf.setFillColor(ink)
                pdf.setFont(font_name, body_size)
                prefix = f"{label}. " if label else ""
                text = f"{prefix}{content}".strip()
                clipped = text if len(text) <= 56 else text[:54] + "..."
                pdf.drawString(x + 10, row_y + 8, clipped)
            y = row_y - 6

    def draw_writing_guides(x: float, top_y: float, available_width: float, area_height: float) -> None:
        """绘制浅色作答横线。"""
        if area_height <= 0:
            return
        guide_color = colors.HexColor("#dbe3ef")
        pdf.setStrokeColor(guide_color)
        pdf.setLineWidth(0.55)
        line_y = top_y - 18
        bottom_y = top_y - area_height
        while line_y > bottom_y + 8:
            pdf.line(x, line_y, x + available_width, line_y)
            line_y -= 18
        pdf.setLineWidth(1)

    def draw_writing_area(area_height: float, *, x: float, available_width: float) -> None:
        """在普通解答题题干后预留作答空间。"""
        nonlocal y
        if area_height <= 0:
            return
        ensure_space(area_height + 12)
        y -= 2
        draw_writing_guides(x, y, available_width, area_height)
        y -= area_height + 8

    def draw_card(lines: list[str], *, label: str = "", sub_type: str = "", score: Any = 0, answer: str = "", analysis: str = "", writing_space_height: float = 0) -> None:
        """绘制小问卡片。"""
        nonlocal y
        content_lines = [line for line in lines if str(line).strip()]
        body_wrapped: list[str] = []
        for line in content_lines:
            body_wrapped.extend(wrap_pdf_line(line, font_name, int(body_size), max_width - 44))
        answer_lines = markdown_lines(answer) if include_answer and answer else []
        analysis_lines = markdown_lines(analysis) if include_answer and analysis else []
        extra_count = len(answer_lines) + len(analysis_lines)
        writing_gap = 10 if writing_space_height else 0
        box_h = 30 + max(1, len(body_wrapped)) * line_height + writing_space_height + writing_gap + extra_count * line_height + (16 if extra_count else 0)
        ensure_space(box_h + 10)
        top = y
        bottom = y - box_h
        pdf.setFillColor(card_fill)
        pdf.setStrokeColor(border)
        pdf.roundRect(margin_x + 14, bottom, max_width - 14, box_h, 9, stroke=1, fill=1)
        x = margin_x + 26
        header_y = top - 17
        pdf.setFillColor(ink)
        pdf.setFont(font_name, body_size)
        pdf.drawString(x, header_y, label)
        next_x = x + 28
        next_x = draw_badge(type_label(sub_type), next_x, header_y - 1)
        pdf.setFillColor(muted)
        pdf.setFont(font_name, small_size)
        pdf.drawString(next_x, header_y, f"{float(score or 0):g} 分")
        y = header_y - 18
        draw_text_lines(body_wrapped, x, max_width - 44, font_size=body_size, color=ink, leading=line_height)
        if writing_space_height:
            y -= 4
            draw_writing_guides(x, y, max_width - 48, writing_space_height)
            y -= writing_space_height + 6
        if answer_lines or analysis_lines:
            y -= 2
            pdf.setStrokeColor(border)
            pdf.line(x, y, margin_x + max_width - 12, y)
            y -= 12
            if answer_lines:
                draw_text_lines(["答案：" + " ".join(answer_lines)], x, max_width - 44, font_size=small_size + 0.5, color=muted, leading=13)
            if analysis_lines:
                draw_text_lines(["解析：" + " ".join(analysis_lines)], x, max_width - 44, font_size=small_size + 0.5, color=muted, leading=13)
        y = bottom - 10

    def draw_question_image(image_path: Path) -> None:
        """执行 draw question image 逻辑。"""
        nonlocal y
        try:
            with Image.open(image_path) as image:
                image_width, image_height = image.size
            if image_width <= 0 or image_height <= 0:
                return
            max_image_width = min(max_width - 56, 310)
            max_image_height = 180
            scale = min(max_image_width / image_width, max_image_height / image_height, 1.0)
            draw_width = image_width * scale
            draw_height = image_height * scale
            ensure_space(draw_height + 14)
            x = margin_x + 24
            y -= 6
            pdf.drawImage(ImageReader(str(image_path)), x, y - draw_height, width=draw_width, height=draw_height, preserveAspectRatio=True, mask="auto")
            y -= draw_height + 12
            pdf.setFont(font_name, body_size)
        except Exception:
            for wrapped in wrap_pdf_line(f"[题图无法导出：{image_path.name}]", font_name, body_size, max_width):
                ensure_space()
                pdf.drawString(margin_x, y, wrapped)
                y -= line_height

    header = paper.get("header") or {}
    total_score = sum(float(question.get("score", 0) or 0) for question in questions)
    title = str(paper.get("title") or "组卷")
    pdf.setFillColor(ink)
    pdf.setFont(font_name, title_size)
    pdf.drawCentredString(width / 2, y, title)
    y -= 23
    meta_parts = []
    if header.get("school"):
        meta_parts.append(f"学校：{header['school']}")
    if header.get("duration"):
        meta_parts.append(f"考试时间：{header['duration']}")
    meta_parts.extend([f"满分：{total_score:g} 分", f"题量：{len(questions)} 题"])
    pdf.setFillColor(muted)
    pdf.setFont(font_name, body_size)
    pdf.drawCentredString(width / 2, y, "　".join(meta_parts))
    y -= 22
    pdf.drawCentredString(width / 2, y, "姓名：____________    班级：____________    考号：____________")
    y -= 18
    pdf.setStrokeColor(colors.HexColor("#475569"))
    pdf.setLineWidth(1.2)
    pdf.line(margin_x, y, width - margin_x, y)
    y -= 24
    if header.get("instructions"):
        draw_text_lines(markdown_lines(header["instructions"]), margin_x, max_width, font_size=small_size + 0.5, color=muted, leading=14)
        y -= 8
    for index, question in enumerate(questions, start=1):
        ensure_space(54)
        row_y = y
        pdf.setFillColor(ink)
        pdf.setFont(font_name, 12)
        pdf.drawString(margin_x, row_y, f"{index}.")
        badge_x = margin_x + 18
        badge_x = draw_badge(type_label(question.get("type")), badge_x, row_y - 1)
        sub_questions = question_sub_questions(question)
        if sub_questions:
            total_sub_count = int(question.get("_subQuestionTotal") or len(sub_questions))
            badge_x = draw_badge(f"已选 {len(sub_questions)}/{total_sub_count} 小问", badge_x, row_y - 1)
        pdf.setFillColor(muted)
        pdf.setFont(font_name, small_size + 0.5)
        pdf.drawString(badge_x + 2, row_y, f"（{float(question.get('score', 0) or 0):g} 分）")
        y -= 22

        raw_markdown = str(question.get("manualMarkdown") or question.get("stemMarkdown") or "")
        stem_markdown, task_options = split_tasks_options(raw_markdown)
        stem_lines = markdown_lines(stem_markdown)
        if stem_lines:
            draw_text_lines(stem_lines, margin_x + 18, max_width - 18, font_size=body_size, color=ink, leading=line_height)
            y -= 3
        for image_path in export_question_image_paths(question, include_sub_questions=False):
            draw_question_image(image_path)
        options = question.get("options") or task_options
        if options:
            draw_options(options)

        for sub_index, sub_question in enumerate(sub_questions, start=1):
            label, sub_stem_markdown, sub_options = sub_question_markdown(sub_question, sub_index)
            sub_type = effective_question_type(sub_question, question.get("type"))
            draw_card(
                markdown_lines(sub_stem_markdown),
                label=label,
                sub_type=sub_type,
                score=sub_question.get("score", 0),
                answer=str(sub_question.get("answer") or ""),
                analysis=str(sub_question.get("analysis") or ""),
                writing_space_height=answer_space_height(sub_question, fallback_type=question.get("type"), is_sub_question=True),
            )
            for image_path in export_question_image_paths(sub_question, include_sub_questions=False):
                draw_question_image(image_path)
            if sub_options:
                draw_options(sub_options)

        if not sub_questions:
            draw_writing_area(
                answer_space_height(question, is_sub_question=False),
                x=margin_x + 18,
                available_width=max_width - 18,
            )

        if include_answer and not sub_questions:
            answer_lines = markdown_lines(question.get("answer")) if question.get("answer") else []
            analysis_lines = markdown_lines(question.get("analysis")) if question.get("analysis") else []
            if answer_lines:
                draw_text_lines(["答案：" + " ".join(answer_lines)], margin_x + 18, max_width - 18, font_size=small_size + 0.5, color=muted, leading=13)
            if analysis_lines:
                draw_text_lines(["解析：" + " ".join(analysis_lines)], margin_x + 18, max_width - 18, font_size=small_size + 0.5, color=muted, leading=13)
        y -= 16
        ensure_space()
    pdf.save()
    return export_path


def strip_markdown_images_for_pandoc(markdown: str) -> str:
    """移除 Pandoc 输入中的 Markdown 图片。"""
    return "\n".join(line for line in str(markdown or "").splitlines() if not re.search(r"!\[[^\]]*\]\([^)]+\)", line)).strip()


def normalize_tasks_for_pandoc(markdown: str) -> str:
    """将 tasks 环境转换为 Pandoc 友好格式。"""
    text = str(markdown or "")

    def replace_tasks(match: re.Match[str]) -> str:
        """执行 replace tasks 逻辑。"""
        task_parts = re.split(r"\\task\b", match.group("body"))[1:]
        lines = [f"- {part.strip()}" for part in task_parts if part.strip()]
        return "\n".join(lines)

    return re.sub(
        r"\\begin\{tasks\}(?:\([^)]+\))?(?P<body>.*?)\\end\{tasks\}",
        replace_tasks,
        text,
        flags=re.S,
    )


def normalize_display_math_for_pandoc(markdown: str) -> str:
    """规范化展示公式，避免空行打断 Pandoc 的 $$ 公式块。"""
    text = str(markdown or "")

    def replace_display(match: re.Match[str]) -> str:
        body = match.group(1).strip()
        body = re.sub(r"\n\s*\n+", "\n", body)
        body = re.sub(r"\\displaystyle\s*", "", body)
        body = re.sub(
            r"\\begin\{array\}\s*\{\s*([^{}]+?)\s*\}",
            lambda item: r"\begin{array}{" + re.sub(r"\s+", "", item.group(1)) + "}",
            body,
        )
        return f"$$\n{body.strip()}\n$$"

    return re.sub(r"(?<!\\)\$\$(.*?)(?<!\\)\$\$", replace_display, text, flags=re.S)


def pandoc_markdown_text(markdown: Any) -> str:
    """规范化传给 Pandoc 的 Markdown 文本。"""
    text = strip_markdown_images_for_pandoc(str(markdown or ""))
    text = normalize_tasks_for_pandoc(text)
    text = normalize_display_math_for_pandoc(text)
    return text.strip()


def pandoc_image_markdown(image_path: Path) -> str:
    """生成 Pandoc 图片 Markdown。"""
    return f"![](<{image_path.resolve().as_posix()}>)"


def latex_escape_text(text: Any) -> str:
    """转义普通文本，数学片段由调用方单独保留。"""
    value = str(text or "")
    replacements = {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
        "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}",
    }
    return "".join(replacements.get(char, char) for char in value)


def latex_clean_markdown_text(text: str) -> str:
    """清理普通 Markdown 标记，保留原始文字。"""
    value = str(text or "")
    value = re.sub(r"`([^`]*)`", r"\1", value)
    value = re.sub(r"\*\*([^*]+)\*\*", r"\1", value)
    value = re.sub(r"^\s{0,3}#{1,6}\s*", "", value, flags=re.M)
    value = re.sub(r"^\s{0,3}>\s?", "", value, flags=re.M)
    value = re.sub(r"^\s*[-*+]\s+", "", value, flags=re.M)
    value = re.sub(r"(?<!\*)\*(?!\*)([^*\n]+)(?<!\*)\*(?!\*)", r"\1", value)
    return value


def latex_clean_math(math: str) -> str:
    """规范化 OCR 产生的 LaTeX 数学片段。"""
    value = str(math or "").strip()
    value = re.sub(r"\n\s*\n+", "\n", value)
    value = re.sub(r"\\displaystyle\s*", "", value)
    value = re.sub(
        r"\\begin\{array\}\s*\{\s*([^{}]+?)\s*\}",
        lambda item: r"\begin{array}{" + re.sub(r"\s+", "", item.group(1)) + "}",
        value,
    )
    value = re.sub(r"\\end\{array\}", r"\\end{array}", value)
    value = re.sub(r"\\begin\s+\{", r"\\begin{", value)
    value = re.sub(r"\\end\s+\{", r"\\end{", value)
    value = re.sub(r"_{2,}", r"\\underline{\\hspace{1.2cm}}", value)
    return value


def latex_text_fragment(text: str) -> str:
    """渲染非数学 Markdown 片段。"""
    cleaned = latex_clean_markdown_text(text)
    lines = [latex_escape_text(line.strip()) for line in cleaned.splitlines() if line.strip()]
    return r"\par ".join(lines)


def latex_inline_markdown_line(line: str) -> str:
    """渲染单行 Markdown，保留行内数学公式。"""
    cleaned = latex_clean_markdown_text(line)
    parts: list[str] = []
    pattern = re.compile(r"(?<!\\)\$(.*?)(?<!\\)\$")
    cursor = 0
    for match in pattern.finditer(cleaned):
        if match.start() > cursor:
            parts.append(latex_escape_text(cleaned[cursor : match.start()]))
        parts.append("$" + latex_clean_math(match.group(1)) + "$")
        cursor = match.end()
    if cursor < len(cleaned):
        parts.append(latex_escape_text(cleaned[cursor:]))
    return "".join(parts).strip()


def latex_paragraph_lines(text: str) -> list[str]:
    """按原始换行生成 LaTeX 段落，避免 PDF 导出把证明步骤合并成一行。"""
    return [rendered for line in str(text or "").splitlines() if line.strip() for rendered in [latex_inline_markdown_line(line)] if rendered]


def latex_markdown_text(markdown: Any) -> str:
    """将 Markdown 转为可由 XeLaTeX 渲染的内容，保留数学公式。"""
    text = strip_markdown_images_for_pandoc(str(markdown or ""))
    text = normalize_tasks_for_pandoc(text)
    parts: list[str] = []
    pattern = re.compile(r"(?<!\\)\$\$(.*?)(?<!\\)\$\$", re.S)
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            parts.extend(latex_paragraph_lines(text[cursor : match.start()]))
        display_math = match.group(1)
        parts.append("\\[\n" + latex_clean_math(display_math) + "\n\\]")
        cursor = match.end()
    if cursor < len(text):
        parts.extend(latex_paragraph_lines(text[cursor:]))
    return "\n\\par\n".join(part for part in parts if part).strip()


def latex_type_label(value: Any) -> str:
    """题型中文标签。"""
    mapping = {
        "choice": "选择题",
        "fill_blank": "填空题",
        "solution": "解答题",
        "unknown": "未知",
    }
    return mapping.get(str(value or "unknown"), str(value or "未知"))


def latex_badge(text: Any) -> str:
    """生成预览风格题型标签。"""
    label = latex_escape_text(str(text or "").strip())
    return rf"\badgebox{{{label}}}" if label else ""


def latex_score_text(value: Any) -> str:
    """生成分值文本。"""
    return rf"\textcolor{{muted}}{{（{numeric_score(value):g} 分）}}"


def latex_answer_lines(area_height: float) -> str:
    """按 ReportLab 留白高度生成 LaTeX 作答横线。"""
    if area_height <= 0:
        return ""
    count = max(3, min(12, int(round(area_height / 18))))
    line = r"\noindent\textcolor{answerline}{\rule{\linewidth}{0.25pt}}\par\vspace{1.05em}"
    return "\n".join(["\\vspace{0.4em}", *([line] * count)])


def latex_image_block(image_path: Path, *, width_ratio: float = 0.42) -> str:
    """生成题图 LaTeX。"""
    path = image_path.resolve().as_posix()
    return rf"\par\vspace{{0.35em}}\includegraphics[width={width_ratio:.2f}\linewidth]{{{path}}}\par\vspace{{0.4em}}"


def latex_options_block(options: list[dict[str, Any]]) -> str:
    """生成两列选项。"""
    items: list[str] = []
    for index, option in enumerate(options):
        if not isinstance(option, dict):
            continue
        label = str(option.get("label") or chr(65 + index)).strip()
        content = latex_markdown_text(option.get("contentMarkdown") or option.get("content") or "")
        if label or content:
            items.append(rf"\textbf{{{latex_escape_text(label)}.}} {content}")
    if not items:
        return ""
    rows: list[str] = [r"\begin{tabularx}{\linewidth}{@{}>{\raggedright\arraybackslash}X >{\raggedright\arraybackslash}X@{}}"]
    for start in range(0, len(items), 2):
        left = items[start]
        right = items[start + 1] if start + 1 < len(items) else ""
        rows.append(left + " & " + right + r"\\[0.6em]")
    rows.append(r"\end{tabularx}")
    return "\n".join(rows)


def latex_answer_block(title: str, content: Any) -> str:
    """生成答案或解析块。"""
    body = latex_markdown_text(content)
    if not body:
        return ""
    return rf"\par\vspace{{0.35em}}\textcolor{{muted}}{{\textbf{{{latex_escape_text(title)}：}} {body}}}"


def latex_question_block(question: dict[str, Any], number: int, include_answer: bool) -> str:
    """生成单题 XeLaTeX 内容。"""
    raw_markdown = str(question.get("manualMarkdown") or question.get("stemMarkdown") or "")
    stem_markdown, task_options = split_tasks_options(raw_markdown)
    question_type = effective_question_type(question)
    sub_questions = question_sub_questions(question)
    badges = [latex_badge(latex_type_label(question_type))]
    if sub_questions:
        total_sub_count = int(question.get("_subQuestionTotal") or len(sub_questions))
        badges.append(latex_badge(f"已选 {len(sub_questions)}/{total_sub_count} 小问"))
    parts: list[str] = [
        r"\vspace{0.8em}",
        rf"\noindent{{\Large\bfseries {number}.}}\quad {' '.join(badges)}\quad {latex_score_text(question.get('score', 0))}",
        r"\par\vspace{0.55em}",
    ]
    stem = latex_markdown_text(stem_markdown)
    if stem:
        parts.append(rf"\noindent {stem}\par")
    for image_path in export_question_image_paths(question, include_sub_questions=False):
        parts.append(latex_image_block(image_path))
    options = question.get("options") or task_options
    if options:
        parts.append(latex_options_block(options))

    for sub_index, sub_question in enumerate(sub_questions, start=1):
        label, sub_stem_markdown, sub_options = sub_question_markdown(sub_question, sub_index)
        sub_type = effective_question_type(sub_question, question.get("type"))
        parts.extend(
            [
                r"\begin{subquestionbox}",
                rf"\textbf{{{latex_escape_text(label)}}}\quad {latex_badge(latex_type_label(sub_type))}\quad \textcolor{{muted}}{{{numeric_score(sub_question.get('score')):g} 分}}\par\vspace{{0.45em}}",
            ]
        )
        sub_stem = latex_markdown_text(sub_stem_markdown)
        if sub_stem:
            parts.append(sub_stem + r"\par")
        for image_path in export_question_image_paths(sub_question, include_sub_questions=False):
            parts.append(latex_image_block(image_path, width_ratio=0.52))
        if sub_options:
            parts.append(latex_options_block(sub_options))
        writing_lines = latex_answer_lines(answer_space_height(sub_question, fallback_type=question.get("type"), is_sub_question=True))
        if writing_lines:
            parts.append(writing_lines)
        if include_answer:
            answer = latex_answer_block("答案", sub_question.get("answer"))
            analysis = latex_answer_block("解析", sub_question.get("analysis"))
            if answer or analysis:
                parts.append(r"\par\textcolor{border}{\rule{\linewidth}{0.35pt}}")
                if answer:
                    parts.append(answer)
                if analysis:
                    parts.append(analysis)
        parts.append(r"\end{subquestionbox}")

    if not sub_questions:
        writing_lines = latex_answer_lines(answer_space_height(question, is_sub_question=False))
        if writing_lines:
            parts.append(writing_lines)
        if include_answer:
            answer = latex_answer_block("答案", question.get("answer"))
            analysis = latex_answer_block("解析", question.get("analysis"))
            if answer:
                parts.append(answer)
            if analysis:
                parts.append(analysis)
    return "\n".join(part for part in parts if part)


def export_paper_pdf_xelatex(paper: dict[str, Any], questions: list[dict[str, Any]], include_answer: bool) -> Path:
    """使用 XeLaTeX 输出预览样式 PDF，保留数学公式排版。"""
    xelatex_command = shutil.which("xelatex")
    if not xelatex_command:
        raise RuntimeError("XeLaTeX is not installed or not available on PATH")
    questions = questions_for_export(paper, questions)
    export_path = EXPORT_ROOT / f"{paper['id']}.pdf"
    header = paper.get("header") or {}
    total_score = sum(numeric_score(question.get("score")) for question in questions)
    meta_parts: list[str] = []
    if header.get("school"):
        meta_parts.append(f"学校：{header['school']}")
    if header.get("duration"):
        meta_parts.append(f"考试时间：{header['duration']}")
    meta_parts.extend([f"满分：{total_score:g} 分", f"题量：{len(questions)} 题"])
    question_blocks = "\n".join(latex_question_block(question, index, include_answer) for index, question in enumerate(questions, start=1))
    instructions = latex_markdown_text(header.get("instructions") or "")
    instructions_line = rf"\par\textcolor{{muted}}{{考生须知：}} {instructions}" if instructions else ""
    tex = rf"""
\documentclass[11pt,a4paper]{{ctexart}}
\usepackage{{geometry}}
\usepackage{{xcolor}}
\usepackage{{amsmath,amssymb}}
\usepackage{{graphicx}}
\usepackage{{tabularx}}
\usepackage{{array}}
\usepackage{{tcolorbox}}
\tcbuselibrary{{breakable}}
\geometry{{left=18mm,right=18mm,top=18mm,bottom=18mm}}
\pagestyle{{empty}}
\setlength{{\parindent}}{{0pt}}
\setlength{{\parskip}}{{0.35em}}
\linespread{{1.12}}\selectfont
\definecolor{{ink}}{{HTML}}{{0F172A}}
\definecolor{{muted}}{{HTML}}{{64748B}}
\definecolor{{border}}{{HTML}}{{DBE3EF}}
\definecolor{{cardfill}}{{HTML}}{{F8FAFC}}
\definecolor{{badgefill}}{{HTML}}{{EFF6FF}}
\definecolor{{badgetext}}{{HTML}}{{0369A1}}
\definecolor{{answerline}}{{HTML}}{{D7E2EF}}
\newtcbox{{\badgebox}}{{on line, colback=badgefill, colframe=blue!25, coltext=badgetext, arc=7pt, boxrule=0.45pt, left=5pt, right=5pt, top=1pt, bottom=1pt, boxsep=0pt}}
\newtcolorbox{{subquestionbox}}{{breakable, colback=cardfill, colframe=border, arc=8pt, boxrule=0.55pt, left=10pt, right=10pt, top=8pt, bottom=8pt, before skip=0.8em, after skip=0.8em}}
\begin{{document}}
\color{{ink}}
\begin{{center}}
{{\LARGE\bfseries {latex_escape_text(paper.get("title") or "组卷")}}}\\[0.75em]
{{\color{{muted}} {latex_escape_text("　".join(meta_parts))}}}\\[0.85em]
{{\color{{muted}} 姓名：\underline{{\hspace{{2.6cm}}}}\quad 班级：\underline{{\hspace{{2.6cm}}}}\quad 考号：\underline{{\hspace{{2.6cm}}}}}}
\end{{center}}
\vspace{{0.45em}}
\noindent\textcolor{{muted}}{{\rule{{\linewidth}}{{0.7pt}}}}
{instructions_line}
{question_blocks}
\end{{document}}
"""
    with tempfile.TemporaryDirectory(prefix="paper_pdf_", dir=str(EXPORT_ROOT)) as temp_dir:
        work_dir = Path(temp_dir)
        tex_path = work_dir / "paper.tex"
        tex_path.write_text(tex, encoding="utf-8")
        completed = subprocess.run(
            [
                xelatex_command,
                "-interaction=nonstopmode",
                "-halt-on-error",
                "-output-directory",
                str(work_dir),
                tex_path.name,
            ],
            cwd=work_dir,
            capture_output=True,
            text=True,
            timeout=120,
        )
        pdf_path = work_dir / "paper.pdf"
        if completed.returncode != 0 or not pdf_path.exists():
            error_path = export_path.with_suffix(".xelatex-error.txt")
            log_path = work_dir / "paper.log"
            error_path.write_text(
                (completed.stdout or "") + "\n" + (completed.stderr or "") + "\n" + (log_path.read_text(encoding="utf-8", errors="ignore") if log_path.exists() else ""),
                encoding="utf-8",
            )
            raise RuntimeError(f"XeLaTeX export failed; see {error_path.name}")
        shutil.copyfile(pdf_path, export_path)
    return export_path


def export_question_markdown(question: dict[str, Any], number: int, include_answer: bool) -> str:
    """生成单题 Pandoc 导出 Markdown。"""
    raw_markdown = str(question.get("manualMarkdown") or question.get("stemMarkdown") or "")
    stem_markdown, task_options = split_tasks_options(raw_markdown)
    lines: list[str] = [f"## {number}. （{float(question.get('score', 0) or 0):g} 分）", ""]
    if stem_markdown.strip():
        lines.extend([pandoc_markdown_text(stem_markdown), ""])
    for image_path in export_question_image_paths(question, include_sub_questions=False):
        lines.extend([pandoc_image_markdown(image_path), ""])

    options = question.get("options") or task_options
    if options:
        for option in options:
            label = str(option.get("label") or "").strip()
            content = pandoc_markdown_text(option.get("content") or "")
            prefix = f"- **{label}.** " if label else "- "
            lines.append(f"{prefix}{content}".rstrip())
        lines.append("")

    sub_questions = question_sub_questions(question)
    for sub_index, sub_question in enumerate(sub_questions, start=1):
        label, sub_stem_markdown, sub_options = sub_question_markdown(sub_question, sub_index)
        sub_stem = pandoc_markdown_text(sub_stem_markdown)
        if sub_stem.startswith("$$"):
            lines.extend([f"**{label}**", "", sub_stem, ""])
        else:
            lines.extend([f"**{label}** {sub_stem}".rstrip(), ""])
        for image_path in export_question_image_paths(sub_question, include_sub_questions=False):
            lines.extend([pandoc_image_markdown(image_path), ""])
        for option in sub_options:
            option_label = str(option.get("label") or "").strip()
            content = pandoc_markdown_text(option.get("content") or "")
            prefix = f"  - **{option_label}.** " if option_label else "  - "
            lines.append(f"{prefix}{content}".rstrip())
        if sub_options:
            lines.append("")
        if include_answer:
            if sub_question.get("answer"):
                lines.extend([f"**{label} 答案：**", "", pandoc_markdown_text(sub_question.get("answer")), ""])
            if sub_question.get("analysis"):
                lines.extend([f"**{label} 解析：**", "", pandoc_markdown_text(sub_question.get("analysis")), ""])

    if include_answer:
        if not sub_questions and question.get("answer"):
            lines.extend(["**答案：**", "", pandoc_markdown_text(question.get("answer")), ""])
        if not sub_questions and question.get("analysis"):
            lines.extend(["**解析：**", "", pandoc_markdown_text(question.get("analysis")), ""])
    return "\n".join(lines).strip()


def export_paper_markdown(paper: dict[str, Any], questions: list[dict[str, Any]], include_answer: bool) -> Path:
    """生成整张试卷的 Pandoc 输入 Markdown 文件。"""
    questions = questions_for_export(paper, questions)
    export_path = EXPORT_ROOT / f"{paper['id']}.md"
    header = paper.get("header") or {}
    total_score = sum(float(question.get("score", 0) or 0) for question in questions)
    lines = [
        f"# {paper.get('title') or '组卷'}",
        "",
    ]
    if header.get("subtitle"):
        lines.extend([f"**{header['subtitle']}**", ""])

    meta_parts: list[str] = []
    if header.get("school"):
        meta_parts.append(f"学校：{header['school']}")
    if header.get("duration"):
        meta_parts.append(f"考试时长：{header['duration']}")
    meta_parts.extend([f"满分：{total_score:g} 分", f"题量：{len(questions)} 题"])
    lines.extend(["　".join(meta_parts), "", "姓名：__________　班级：__________　考号：__________", ""])

    if header.get("instructions"):
        lines.extend(["## 考生须知", "", pandoc_markdown_text(header["instructions"]), ""])

    for index, question in enumerate(questions, start=1):
        lines.extend([export_question_markdown(question, index, include_answer), ""])

    export_path.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")
    return export_path


def preferred_pandoc_cjk_font() -> str:
    """选择 Pandoc PDF 导出优先使用的中文字体。"""
    configured = os.getenv("PANDOC_CJK_FONT")
    if configured:
        return configured
    for font_name in ("Songti SC", "PingFang SC", "Heiti SC", "STSong", "Noto Serif CJK SC"):
        return font_name
    return "Songti SC"


def run_pandoc(markdown_path: Path, output_path: Path, output_format: str) -> None:
    """调用 Pandoc 生成指定格式文件。"""
    pandoc_command = shutil.which("pandoc")
    if not pandoc_command:
        raise RuntimeError("Pandoc is not installed or not available on PATH")

    command = [
        pandoc_command,
        str(markdown_path),
        "-f",
        "markdown+tex_math_dollars+tex_math_single_backslash",
        "--standalone",
        "-o",
        str(output_path),
    ]
    if output_format == "pdf":
        xelatex_command = shutil.which("xelatex")
        if not xelatex_command:
            raise RuntimeError("XeLaTeX is not installed or not available on PATH")
        command.extend(
            [
                "--pdf-engine",
                xelatex_command,
                "-V",
                f"CJKmainfont={preferred_pandoc_cjk_font()}",
                "-V",
                "geometry:margin=2cm",
            ]
        )

    completed = subprocess.run(command, cwd=PROJECT_ROOT, capture_output=True, text=True, timeout=120)
    if completed.returncode != 0:
        error_path = output_path.with_suffix(f".pandoc-error.txt")
        error_path.write_text((completed.stdout or "") + "\n" + (completed.stderr or ""), encoding="utf-8")
        raise RuntimeError(f"Pandoc export failed; see {error_path.name}")


def export_flow_status() -> dict[str, Any]:
    """返回导出能力运行时状态。"""
    pandoc_command = shutil.which("pandoc")
    xelatex_command = shutil.which("xelatex")
    return {
        "capability": "export-flow",
        "formats": ["markdown", "docx", "pdf"],
        "strategy": "docx-pandoc-pdf-xelatex-preview-reportlab-fallback",
        "pandocInstalled": pandoc_command is not None,
        "pandocCommand": pandoc_command,
        "xelatexInstalled": xelatex_command is not None,
        "xelatexCommand": xelatex_command,
        "cjkFont": preferred_pandoc_cjk_font(),
        "fallbackEnabled": True,
        "exportRoot": str(EXPORT_ROOT),
    }


def export_paper_docx(paper: dict[str, Any], questions: list[dict[str, Any]], include_answer: bool) -> Path:
    """使用 Pandoc 导出 DOCX。"""
    questions = questions_for_export(paper, questions)
    output_path = EXPORT_ROOT / f"{paper['id']}.docx"
    try:
        markdown_path = export_paper_markdown(paper, questions, include_answer)
        run_pandoc(markdown_path, output_path, "docx")
        return output_path
    except Exception:
        return export_paper_docx_legacy(paper, questions, include_answer)


def export_paper_pdf(paper: dict[str, Any], questions: list[dict[str, Any]], include_answer: bool) -> Path:
    """按前端预览样式导出 PDF。"""
    try:
        return export_paper_pdf_xelatex(paper, questions, include_answer)
    except Exception:
        return export_paper_pdf_legacy(paper, questions, include_answer)
